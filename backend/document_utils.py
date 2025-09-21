from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO

def set_cell_background(cell, fill):
    """
    Set cell background color
    """
    cell_properties = cell._element.tcPr
    if cell_properties is None:
        cell_properties = OxmlElement('w:tcPr')
        cell._element.append(cell_properties)
    
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), fill)
    cell_properties.append(shading)

def create_diet_document(paziente_data, dieta_data):
    """
    Create a Word document containing the patient's diet plan.
    
    Args:
        paziente_data: Dictionary containing patient data
        dieta_data: Dictionary containing diet data
    
    Returns:
        BytesIO object containing the Word document
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import RGBColor, Inches
    from docx.enum.section import WD_ORIENT
    
    def set_cell_background(cell, color_hex):
        """Set background color for table cell"""
        cell_xml_element = cell._tc
        table_cell_properties = cell_xml_element.get_or_add_tcPr()
        
        # Remove any existing shading first
        for shd in table_cell_properties.xpath('.//w:shd'):
            table_cell_properties.remove(shd)
            
        # Add new shading
        shade_obj = OxmlElement('w:shd')
        shade_obj.set(qn('w:fill'), color_hex)
        shade_obj.set(qn('w:val'), 'clear')
        table_cell_properties.append(shade_obj)
    
    # Create a new document
    doc = Document()
    
    # Set document to horizontal/landscape orientation
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    # Swap width and height for landscape
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    
    # Set narrow margins
    from docx.shared import Inches
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    
    # Set default font for the document to Century Gothic
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Century Gothic'
    font.size = Pt(12)
    
    # Set document title with specific styling
    title = doc.add_paragraph()  # Use paragraph instead of heading to remove delimiter
    title_run = title.add_run(f'Piano nutrizionale {paziente_data.get("nome", "Paziente")}')
    title_run.font.name = 'Muthiara -Demo Version-'
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(0x77, 0x20, 0x6d)  # #77206d
    title_run.bold = False  # Regular style
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)  # Reduced spacing after title
    
    # Add subtitle with same font and new color
    subtitle = doc.add_paragraph()  # Use paragraph instead of heading to remove delimiter
    subtitle_run = subtitle.add_run('Piano a scelta libera')
    subtitle_run.font.name = 'Muthiara -Demo Version-'
    subtitle_run.font.size = Pt(20)
    subtitle_run.font.color.rgb = RGBColor(0xe5, 0x9e, 0xdc)  # #e59edc
    subtitle_run.bold = False  # Regular style
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(3)  # Reduced spacing before subtitle
    
    # Add patient information if needed (without heading)
    if any(paziente_data.get(key) for key in ['cognome', 'eta', 'email', 'telefono']):
        patient_info = doc.add_paragraph()
        
        if paziente_data.get('cognome'):
            patient_run = patient_info.add_run(f"Nome: {paziente_data['nome']} {paziente_data['cognome']}")
            patient_run.bold = True
            patient_run.font.name = 'Century Gothic'
            patient_run.font.size = Pt(12)
            patient_run.font.color.rgb = RGBColor(0, 0, 0)  # Black
        
        if paziente_data.get('eta'):
            age_run = patient_info.add_run(f"\nEtà: {paziente_data['eta']} anni")
            age_run.font.name = 'Century Gothic'
            age_run.font.size = Pt(12)
            age_run.font.color.rgb = RGBColor(0, 0, 0)  # Black
    
    # Define meal titles in UPPERCASE
    meal_titles = {
        'colazione': 'COLAZIONE',
        'spuntino': 'SPUNTINO',
        'pranzo': 'PRANZO',
        'merenda': 'MERENDA',
        'cena': 'CENA'
    }
    
    def categorize_food_by_nutrition(alimento):
        """
        Categorize food by nutritional type based on food name or category.
        Returns one of: 'FONTI DI CARBOIDRATI', 'FONTI DI PROTEINE', 'FONTI DI GRASSI', 'CONTORNI'
        """
        nome = alimento.get('nome', '').lower()
        categoria = alimento.get('categoria', '').lower()
        tipo = alimento.get('tipo', '').lower()
        
        # Carboidrate sources
        carb_keywords = ['pane', 'pasta', 'riso', 'cereali', 'farro', 'orzo', 'avena', 
                        'quinoa', 'patate', 'patata', 'biscotti', 'crackers', 'fette', 
                        'gallette', 'muesli', 'cornflakes', 'fiocchi']
        
        # Protein sources
        protein_keywords = ['carne', 'pollo', 'manzo', 'maiale', 'pesce', 'salmone', 
                           'tonno', 'merluzzo', 'uova', 'uovo', 'formaggio', 'ricotta', 
                           'mozzarella', 'parmigiano', 'legumi', 'fagioli', 'lenticchie', 
                           'ceci', 'piselli', 'tofu', 'seitan', 'prosciutto', 'bresaola']
        
        # Fat sources
        fat_keywords = ['olio', 'burro', 'noci', 'mandorle', 'nocciole', 'semi', 
                       'avocado', 'olive', 'oliva']
        
        # Vegetables/sides
        vegetable_keywords = ['verdura', 'insalata', 'spinaci', 'broccoli', 'zucchine', 
                             'pomodori', 'carote', 'peperoni', 'melanzane', 'contorno']
        
        # Check against keywords
        text_to_check = f"{nome} {categoria} {tipo}"
        
        if any(keyword in text_to_check for keyword in carb_keywords):
            return 'FONTI DI CARBOIDRATI'
        elif any(keyword in text_to_check for keyword in protein_keywords):
            return 'FONTI DI PROTEINE'
        elif any(keyword in text_to_check for keyword in fat_keywords):
            return 'FONTI DI GRASSI'
        elif any(keyword in text_to_check for keyword in vegetable_keywords):
            return 'CONTORNI'
        else:
            # Default categorization based on original category if available
            if 'carboidrat' in text_to_check or 'cereale' in text_to_check:
                return 'FONTI DI CARBOIDRATI'
            elif 'protein' in text_to_check or 'carne' in text_to_check:
                return 'FONTI DI PROTEINE'
            elif 'grass' in text_to_check or 'condimento' in text_to_check:
                return 'FONTI DI GRASSI'
            elif 'verdur' in text_to_check or 'contorn' in text_to_check:
                return 'CONTORNI'
            else:
                return 'FONTI DI CARBOIDRATI'  # Default fallback
    
    def should_use_table(meal_data):
        """
        Determine if meal should use table format based on:
        1. Multiple categories (different food types) - based on main foods only
        2. Any food item has more than 2 equivalents
        """
        if not meal_data or not meal_data.get('alimenti'):
            return False
        
        # Check for multiple nutritional categories based on main foods only
        categories = set()
        for alimento in meal_data['alimenti']:
            nutrition_category = categorize_food_by_nutrition(alimento)
            categories.add(nutrition_category)
        
        # Use table if more than 1 nutritional category
        if len(categories) > 1:
            return True
        
        # Use table if any food item has more than 2 equivalents
        for alimento in meal_data['alimenti']:
            if alimento.get('equivalenti') and len(alimento['equivalenti']) > 2:
                return True
        
        return False
    
    def get_food_categories_with_equivalents(meal_data):
        """
        Extract and organize food by nutritional categories.
        Only the main food (alimento principale) determines the category.
        All equivalents stay in the same column as their main food.
        """
        if not meal_data or not meal_data.get('alimenti'):
            return {}
        
        # Check if we need special handling for equivalents
        has_many_equivalents = any(
            alimento.get('equivalenti') and len(alimento['equivalenti']) > 2
            for alimento in meal_data['alimenti']
        )
        
        if has_many_equivalents:
            # Create columns based on nutritional categories of main foods only
            categories = {
                'FONTI DI CARBOIDRATI': [],
                'FONTI DI PROTEINE': [],
                'FONTI DI GRASSI': [],
                'CONTORNI': []
            }
            
            for alimento in meal_data['alimenti']:
                # Categorize based only on the main food
                nutrition_category = categorize_food_by_nutrition(alimento)
                
                # Add main food to appropriate category
                categories[nutrition_category].append(alimento)
                
                # Add ALL equivalents to the SAME category as the main food
                if alimento.get('equivalenti'):
                    for equiv in alimento['equivalenti']:
                        categories[nutrition_category].append(equiv)
            
            # Remove empty categories
            return {k: v for k, v in categories.items() if v}
        else:
            # Standard category grouping by nutrition type (main foods only)
            categories = {}
            for alimento in meal_data['alimenti']:
                nutrition_category = categorize_food_by_nutrition(alimento)
                if nutrition_category not in categories:
                    categories[nutrition_category] = []
                categories[nutrition_category].append(alimento)
            
            return categories
    
    def get_food_for_bullets(meal_data):
        """
        Get food organized specifically for bullet point combinations.
        Each main food item becomes its own "category" for combination purposes.
        """
        if not meal_data or not meal_data.get('alimenti'):
            return {}
        
        # Each main food item becomes its own category for combinations
        categories = {}
        for i, alimento in enumerate(meal_data['alimenti']):
            # Use food name or index as category to ensure separate grouping
            category_key = f"FOOD_{i+1}_{alimento['nome'][:20]}"  # Truncate long names
            categories[category_key] = [alimento]
        
        return categories
    
    def format_food_with_quantity(alimento):
        """Format food item with quantity - NOT uppercase"""
        name = alimento['nome']  # Remove .upper()
        if alimento.get('quantita') and alimento.get('unita'):
            quantity = alimento['quantita']
            unit = alimento['unita']
            
            if unit.lower() in ['g', 'grammi']:
                if quantity == int(quantity):
                    return f"{name} ({int(quantity)} g)"
                else:
                    return f"{name} ({quantity} g)"
            else:
                if quantity == int(quantity):
                    return f"{name} ({int(quantity)} {unit})"
                else:
                    return f"{name} ({quantity} {unit})"
        return name
    
    def create_bullet_combinations(categoria_foods):
        """
        Create bullet point combinations for simple meals.
        Each main food + equivalents should be combined with other foods.
        This creates ALL possible combinations between each food group.
        """
        if not categoria_foods:
            return []
        
        # First, expand each category to include main foods + their equivalents
        expanded_categories = {}
        
        for category_name, foods in categoria_foods.items():
            expanded_foods = []
            for alimento in foods:
                # Add the main food
                expanded_foods.append(alimento)
                # Add all its equivalents
                if alimento.get('equivalenti'):
                    expanded_foods.extend(alimento['equivalenti'])
            expanded_categories[category_name] = expanded_foods
        
        category_names = list(expanded_categories.keys())
        combinations = []
        
        if len(category_names) == 1:
            # Single category - list all items (main + equivalents) separately
            for alimento in expanded_categories[category_names[0]]:
                combinations.append([format_food_with_quantity(alimento)])
        
        elif len(category_names) == 2:
            # Two categories - create ALL combinations between them
            cat1_foods = expanded_categories[category_names[0]]
            cat2_foods = expanded_categories[category_names[1]]
            
            for food1 in cat1_foods:
                for food2 in cat2_foods:
                    combination = [
                        format_food_with_quantity(food1),
                        format_food_with_quantity(food2)
                    ]
                    combinations.append(combination)
        
        else:
            # More than 2 categories - create combinations with Cartesian product
            import itertools
            
            # Get all expanded food lists
            food_lists = [expanded_categories[cat] for cat in category_names]
            
            # Create all combinations using Cartesian product
            for combination_tuple in itertools.product(*food_lists):
                combination = [format_food_with_quantity(food) for food in combination_tuple]
                combinations.append(combination)
        
        return combinations
    
    # Process each meal
    for meal_key, meal_title in meal_titles.items():
        meal_data = dieta_data.get(meal_key)
        if not meal_data or not meal_data.get('alimenti'):
            continue
        
        # Add meal title with "scegli:" or detailed instructions for PRANZO/CENA
        meal_paragraph = doc.add_paragraph()
        
        if meal_key in ['pranzo', 'cena']:
            # Special formatting for PRANZO and CENA with detailed instructions
            meal_run = meal_paragraph.add_run(f"{meal_title}, scegli ")
            meal_run.font.name = 'Century Gothic'
            meal_run.font.size = Pt(14)
            meal_run.font.color.rgb = RGBColor(0, 0, 0)  # Black
            meal_run.bold = True
            
            # Add bold parts for nutritional categories
            carb_run = meal_paragraph.add_run("1 fonte di carboidrati")
            carb_run.font.name = 'Century Gothic'
            carb_run.font.size = Pt(14)
            carb_run.font.color.rgb = RGBColor(0, 0, 0)  # Black
            carb_run.bold = True
            
            comma1_run = meal_paragraph.add_run(", ")
            comma1_run.font.name = 'Century Gothic'
            comma1_run.font.size = Pt(14)
            comma1_run.font.color.rgb = RGBColor(0, 0, 0)  # Black
            comma1_run.bold = True
            
            protein_run = meal_paragraph.add_run("1 fonte di proteine")
            protein_run.font.name = 'Century Gothic'
            protein_run.font.size = Pt(14)
            protein_run.font.color.rgb = RGBColor(0, 0, 0)  # Black
            protein_run.bold = True
            
            comma2_run = meal_paragraph.add_run(", ")
            comma2_run.font.name = 'Century Gothic'
            comma2_run.font.size = Pt(14)
            comma2_run.font.color.rgb = RGBColor(0, 0, 0)  # Black
            comma2_run.bold = True
            
            fat_run = meal_paragraph.add_run("1 fonte di grassi")
            fat_run.font.name = 'Century Gothic'
            fat_run.font.size = Pt(14)
            fat_run.font.color.rgb = RGBColor(0, 0, 0)  # Black
            fat_run.bold = True
            
            and_run = meal_paragraph.add_run(" e ")
            and_run.font.name = 'Century Gothic'
            and_run.font.size = Pt(14)
            and_run.font.color.rgb = RGBColor(0, 0, 0)  # Black
            and_run.bold = True
            
            contorno_run = meal_paragraph.add_run("1 contorno")
            contorno_run.font.name = 'Century Gothic'
            contorno_run.font.size = Pt(14)
            contorno_run.font.color.rgb = RGBColor(0, 0, 0)  # Black
            contorno_run.bold = True
            
            final_run = meal_paragraph.add_run(" per comporre il tuo piatto:")
            final_run.font.name = 'Century Gothic'
            final_run.font.size = Pt(14)
            final_run.font.color.rgb = RGBColor(0, 0, 0)  # Black
            final_run.bold = True
        else:
            # Regular "scegli:" for other meals (not bold)
            meal_run = meal_paragraph.add_run(f"{meal_title}, scegli:")
            meal_run.font.name = 'Century Gothic'
            meal_run.font.size = Pt(14)
            meal_run.font.color.rgb = RGBColor(0, 0, 0)  # Black
            meal_run.bold = True
        
        categories = get_food_categories_with_equivalents(meal_data)
        
        if should_use_table(meal_data):
            # Create table with nutritional categories as columns
            category_names = list(categories.keys())
            table = doc.add_table(rows=1, cols=len(category_names))
            
            # Remove table borders except for header row and vertical lines
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            
            tbl = table._tbl
            tblPr = tbl.tblPr
            
            # Keep all borders including top
            tblBorders = OxmlElement('w:tblBorders')
            
            # Keep all borders (top, left, right, bottom, vertical)
            top = OxmlElement('w:top')
            top.set(qn('w:val'), 'single')
            top.set(qn('w:sz'), '4')
            top.set(qn('w:color'), '000000')
            tblBorders.append(top)
            
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), '4')
            left.set(qn('w:color'), '000000')
            tblBorders.append(left)
            
            right = OxmlElement('w:right')
            right.set(qn('w:val'), 'single')
            right.set(qn('w:sz'), '4')
            right.set(qn('w:color'), '000000')
            tblBorders.append(right)
            
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:color'), '000000')
            tblBorders.append(bottom)
            
            insideV = OxmlElement('w:insideV')
            insideV.set(qn('w:val'), 'single')
            insideV.set(qn('w:sz'), '4')
            insideV.set(qn('w:color'), '000000')
            tblBorders.append(insideV)
            
            # Remove only inside horizontal borders
            insideH = OxmlElement('w:insideH')
            insideH.set(qn('w:val'), 'none')
            tblBorders.append(insideH)
                
            tblPr.append(tblBorders)
            
            # Set headers with bottom border only
            header_cells = table.rows[0].cells
            for i, category_name in enumerate(category_names):
                header_cell = header_cells[i]
                header_cell.text = category_name
                
                # Add bottom border to header cells only
                tc = header_cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '4')
                bottom.set(qn('w:color'), '000000')
                tcBorders.append(bottom)
                tcPr.append(tcBorders)
                
                # Set white background for all header cells
                set_cell_background(header_cell, 'FFFFFF')
                    
                for paragraph in header_cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.space_after = Pt(6)
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.name = 'Century Gothic'
                        run.font.size = Pt(12)
                        run.font.color.rgb = RGBColor(0, 0, 0)  # Black
            
            # Find max items in any category
            max_items = max(len(foods) for foods in categories.values())
            
            # Add rows for food items with extra spacing
            for row_idx in range(max_items):
                row_cells = table.add_row().cells
                
                for col_idx, category_name in enumerate(category_names):
                    foods_in_category = categories[category_name]
                    
                    if row_idx < len(foods_in_category):
                        alimento = foods_in_category[row_idx]
                        cell = row_cells[col_idx]
                        
                        # Clear cell and add formatted content
                        cell.paragraphs[0].clear()
                        paragraph = cell.paragraphs[0]
                        
                        # Add extra spacing between rows
                        paragraph.paragraph_format.space_after = Pt(12)  # Increased spacing
                        paragraph.paragraph_format.space_before = Pt(6)   # Space before
                        
                        # Add bullet and formatted food item
                        bullet_run = paragraph.add_run("- ")
                        bullet_run.font.name = 'Century Gothic'
                        bullet_run.font.size = Pt(12)
                        bullet_run.font.color.rgb = RGBColor(0, 0, 0)  # Black
                        
                        food_run = paragraph.add_run(format_food_with_quantity(alimento))
                        food_run.font.name = 'Century Gothic'
                        food_run.font.size = Pt(12)
                        food_run.font.color.rgb = RGBColor(0, 0, 0)  # Black
        else:
            # Use bullet points with combinations
            bullet_categories = get_food_for_bullets(meal_data)
            combinations = create_bullet_combinations(bullet_categories)
            
            for combination in combinations:
                paragraph = doc.add_paragraph(style='List Bullet')
                paragraph.paragraph_format.space_after = Pt(6)
                
                # Join combination with " + "
                combination_text = " + ".join(combination)
                
                run = paragraph.add_run(combination_text)
                run.font.name = 'Century Gothic'
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)  # Black
        
        # Add meal notes if available
        if meal_data.get('note'):
            doc.add_paragraph()  # Add space
            note_paragraph = doc.add_paragraph()
            note_paragraph.paragraph_format.space_after = Pt(6)
            
            note_run = note_paragraph.add_run(f"Consigli: {meal_data['note']}")
            note_run.italic = True
            note_run.font.name = 'Century Gothic'
            note_run.font.size = Pt(12)
            note_run.font.color.rgb = RGBColor(0, 0, 0)  # Black
        
        # Add spacing between meals
        doc.add_paragraph()
    
    # Add general notes from diet data (without heading)
    if dieta_data.get('note'):
        note_paragraph = doc.add_paragraph()
        note_heading_run = note_paragraph.add_run('Note')
        note_heading_run.font.name = 'Century Gothic'
        note_heading_run.font.size = Pt(12)
        note_heading_run.font.color.rgb = RGBColor(0, 0, 0)  # Black
        note_heading_run.bold = True
        
        # Split notes by line breaks and format properly
        note_lines = dieta_data['note'].split('\n')
        for line in note_lines:
            if line.strip():
                note_paragraph = doc.add_paragraph()
                note_paragraph.paragraph_format.space_after = Pt(6)
                
                # Handle bold formatting within notes
                if '**' in line:
                    parts = line.split('**')
                    for i, part in enumerate(parts):
                        run = note_paragraph.add_run(part)
                        run.font.name = 'Century Gothic'
                        run.font.size = Pt(12)
                        run.font.color.rgb = RGBColor(0, 0, 0)  # Black
                        if i % 2 == 1:  # Odd indices are between ** markers
                            run.bold = True
                else:
                    run = note_paragraph.add_run(line)
                    run.font.name = 'Century Gothic'
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Black
    
    # Add general advice section (without heading style)
    advice_paragraph = doc.add_paragraph()
    advice_heading_run = advice_paragraph.add_run('Alcuni consigli generali')
    advice_heading_run.font.name = 'Muthiara -Demo Version-'
    advice_heading_run.font.size = Pt(18)  # Size 18 as requested
    advice_heading_run.font.color.rgb = RGBColor(0xe5, 0x9e, 0xdc)  # #e59edc (same as subtitle)
    advice_heading_run.bold = False
    
    # Get general tips from diet data or use defaults
    general_tips = dieta_data.get('consigli_generali', [
        "Non è necessario rispettare quantità precise per la **verdura**. Anzi aumentane le quantità durante i pasti se hai ancora fame.",
        "Il **caffè** non è inserito nel piano, ma puoi berne quanto ne vuoi durante la giornata.",
        "Per la **frutta** scegli quella che preferisci (considera 150 g come riferimento) di base corrisponde ad 1 frutto grande (pesca) o 2 piccole (ad esempio le albicocche).",
        "**Spuntini e merende** possono essere **scambiati** tra mattina e pomeriggio, così come **pranzi e cene** sia all'interno della stessa, che tra giornate diverse.",
        "Tutte le cose, pesale la prima settimana, **poi vai ad occhio**!",
        "**L'olio** indicato include sia quello per condire i piatti sia quello usato per la cottura delle verdure, fai attenzione!",
        "Ridurre al **minimo l'aggiunta di sale**.",
        "Preferire **pane e pasta integrali**.",
        "**Evita di trascorrere troppo tempo a digiuno**, ne risentirà il tuo pasto successivo poi!",
        "Bere almeno **2 L di acqua** al giorno (molto importante!!!)"
    ])
    
    for tip in general_tips:
        paragraph = doc.add_paragraph(style='List Bullet')
        paragraph.paragraph_format.space_after = Pt(6)
        
        # Handle bold formatting with ** markers
        parts = tip.split('**')
        for i, part in enumerate(parts):
            run = paragraph.add_run(part)
            run.font.name = 'Century Gothic'
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)  # Black
            if i % 2 == 1:  # Odd indices are between ** markers
                run.bold = True
    
    # Add extra notes if available
    if dieta_data.get('note_extra'):
        doc.add_paragraph()
        extra_paragraph = doc.add_paragraph(dieta_data['note_extra'])
        extra_paragraph.paragraph_format.space_after = Pt(6)
        for run in extra_paragraph.runs:
            run.font.name = 'Century Gothic'
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)  # Black
    
    # Save document to BytesIO
    doc_stream = BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    
    return doc_stream